"""Shared docx/pptx helpers for the bootcamp generators."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x1E, 0x27, 0x61)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
LIGHT_GRAY_HEX = "E6E6E6"
ACCENT_HEX = "1E2761"


def new_doc():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)
    return doc


def add_title(doc, text, color=NAVY, size=22):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = color
    run.font.size = Pt(size)
    return p


def add_subtitle(doc, text, color=DARK_GRAY, size=14):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.color.rgb = color
    run.font.size = Pt(size)
    return p


def add_h1(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = NAVY
    run.font.size = Pt(16)
    return p


def add_h2(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = NAVY
    run.font.size = Pt(13)
    return p


def add_h3(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = DARK_GRAY
    run.font.size = Pt(12)
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.color.rgb = DARK_GRAY
    run.font.size = Pt(11)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    if level:
        pf = p.paragraph_format
        pf.left_indent = Inches(0.25 * (level + 1))
    run = p.add_run(text)
    run.font.color.rgb = DARK_GRAY
    run.font.size = Pt(11)
    return p


def add_numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    run = p.add_run(text)
    run.font.color.rgb = DARK_GRAY
    run.font.size = Pt(11)
    return p


def add_code(doc, code):
    for line in code.split("\n"):
        p = doc.add_paragraph()
        run = p.add_run(line if line else " ")
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x10, 0x10, 0x10)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)


def add_table(doc, header, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(header):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _shade(hdr[i], ACCENT_HEX)
    for r, row in enumerate(rows, start=1):
        cells = table.rows[r].cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
    if col_widths:
        for i, w in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = Inches(w)
    return table


def _shade(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_divider(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), ACCENT_HEX)
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_page_break(doc):
    doc.add_page_break()


def add_header_block(doc, course, doc_type, week_num=None, week_title=None):
    add_title(doc, "Proxiant Academy")
    add_subtitle(doc, course)
    if week_num:
        add_subtitle(doc, f"{doc_type} | Week {week_num}: {week_title}", size=12)
    else:
        add_subtitle(doc, doc_type, size=12)
    add_divider(doc)
