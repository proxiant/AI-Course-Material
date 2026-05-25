"""Textbook DOCX helpers with pedagogical formatting."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x1E, 0x27, 0x61)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
ACCENT = RGBColor(0x2C, 0x5F, 0xF5)
ACCENT_HEX = "1E2761"


def new_doc(book_title=None):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    for section in doc.sections:
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
    if book_title:
        add_page_numbers(doc)
        add_running_header(doc, book_title)
    return doc


def add_page_numbers(doc):
    """Add automatic page numbers to footer (centered)."""
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run.font.size = Pt(10)
    run.font.color.rgb = DARK_GRAY


def add_running_header(doc, title):
    """Add running header with book title."""
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(title)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = DARK_GRAY


def cover_page(doc, title, subtitle, author, publisher, edition, year, isbn=None):
    """Render a publishing-grade cover page sized to fit a single page."""
    for _ in range(2):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(publisher.upper())
    run.bold = True
    run.font.color.rgb = ACCENT
    run.font.size = Pt(14)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Professional Reference Series")
    run.italic = True
    run.font.color.rgb = DARK_GRAY
    run.font.size = Pt(11)
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.color.rgb = NAVY
    run.font.size = Pt(28)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(subtitle)
    run.italic = True
    run.font.color.rgb = DARK_GRAY
    run.font.size = Pt(14)
    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(author)
    run.bold = True
    run.font.color.rgb = NAVY
    run.font.size = Pt(16)
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{edition}  |  {year}")
    run.font.color.rgb = DARK_GRAY
    run.font.size = Pt(11)


def copyright_page(doc, title, author, publisher, edition, year,
                   isbn=None, license_text=None):
    """Render a copyright page."""
    for _ in range(8):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.color.rgb = NAVY
    run.font.size = Pt(13)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"by {author}")
    run.font.color.rgb = DARK_GRAY
    run.font.size = Pt(11)
    doc.add_paragraph()
    doc.add_paragraph()
    items = [
        f"Published by {publisher}, San Francisco Bay Area.",
        f"{edition}, {year}.",
        f"Copyright © {year} {publisher}. All rights reserved.",
        "",
        "No part of this work may be reproduced or transmitted in any form or by "
        "any means without the prior written permission of the publisher, except "
        "for brief quotations in critical reviews and certain other non-commercial "
        "uses permitted by copyright law.",
        "",
        "While every precaution has been taken in the preparation of this book, "
        "the publisher and author assume no responsibility for errors or "
        "omissions, or for damages resulting from the use of the information "
        "contained herein.",
        "",
        "Trademarked product names mentioned in this book are the property of "
        "their respective owners. The use of such names is for editorial purposes "
        "only and does not imply endorsement.",
    ]
    if isbn:
        items.insert(3, f"Reference Identifier (RID): {isbn}")
    for item in items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if item:
            run = p.add_run(item)
            run.font.color.rgb = DARK_GRAY
            run.font.size = Pt(10)


def about_author(doc, author, bio):
    """About-the-author section."""
    p = doc.add_paragraph()
    run = p.add_run("About the Author")
    run.bold = True
    run.font.color.rgb = NAVY
    run.font.size = Pt(22)
    add_divider(doc)
    p = doc.add_paragraph()
    run = p.add_run(author)
    run.bold = True
    run.font.color.rgb = NAVY
    run.font.size = Pt(14)
    p = doc.add_paragraph()
    run = p.add_run(bio)
    run.font.color.rgb = DARK_GRAY
    run.font.size = Pt(11)


def add_title(doc, text, color=NAVY, size=24, center=True):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = color
    run.font.size = Pt(size)
    return p


def add_subtitle(doc, text, color=DARK_GRAY, size=14, center=True, italic=True):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = italic
    run.font.color.rgb = color
    run.font.size = Pt(size)
    return p


def chapter_heading(doc, chapter_label, title):
    """Major chapter heading on a new page."""
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run1 = p1.add_run(chapter_label)
    run1.italic = True
    run1.font.color.rgb = ACCENT
    run1.font.size = Pt(14)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run2 = p2.add_run(title)
    run2.bold = True
    run2.font.color.rgb = NAVY
    run2.font.size = Pt(28)
    add_divider(doc)


def section_heading(doc, number, title):
    p = doc.add_paragraph()
    run = p.add_run(f"{number} {title}")
    run.bold = True
    run.font.color.rgb = NAVY
    run.font.size = Pt(18)


def subsection_heading(doc, number, title):
    p = doc.add_paragraph()
    run = p.add_run(f"{number} {title}")
    run.bold = True
    run.font.color.rgb = NAVY
    run.font.size = Pt(14)


def subsubsection_heading(doc, title):
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.color.rgb = DARK_GRAY
    run.font.size = Pt(12)


def body(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.color.rgb = DARK_GRAY
    run.font.size = Pt(11)
    return p


def callout(doc, label, text, color=ACCENT):
    """Boxed pedagogical callout (Definition, Example, Theorem, etc.)."""
    p = doc.add_paragraph()
    run1 = p.add_run(f"[{label}] ")
    run1.bold = True
    run1.font.color.rgb = color
    run1.font.size = Pt(11)
    run2 = p.add_run(text)
    run2.italic = True
    run2.font.color.rgb = DARK_GRAY
    run2.font.size = Pt(11)


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.color.rgb = DARK_GRAY
    run.font.size = Pt(11)


def numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    run = p.add_run(text)
    run.font.color.rgb = DARK_GRAY
    run.font.size = Pt(11)


def code(doc, text):
    for line in text.split("\n"):
        p = doc.add_paragraph()
        run = p.add_run(line if line else " ")
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)


def equation(doc, text, label=None):
    """Display an equation centered with optional label."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.color.rgb = NAVY
    run.font.size = Pt(12)
    if label:
        run2 = p.add_run(f"     ({label})")
        run2.font.color.rgb = DARK_GRAY
        run2.font.size = Pt(10)


def image(doc, path, caption=None, width=5.5):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=Inches(width))
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        crun = cap.add_run(caption)
        crun.italic = True
        crun.font.color.rgb = DARK_GRAY
        crun.font.size = Pt(10)


def add_divider(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), ACCENT_HEX)
    pBdr.append(bottom)
    pPr.append(pBdr)


def page_break(doc):
    doc.add_page_break()
