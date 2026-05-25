"""Document helpers for the textbook."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x1E, 0x27, 0x61)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
ACCENT = RGBColor(0x2C, 0x5F, 0xF5)
ACCENT_HEX = "1E2761"


def new_doc(book_title=None):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    if book_title:
        _add_page_numbers(doc)
        _add_running_header(doc, book_title)
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)
    return doc


def add_title(doc, text, color=NAVY, size=22, center=False):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = color
    run.font.size = Pt(size)
    return p


def add_subtitle(doc, text, color=DARK_GRAY, size=14, center=False):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.color.rgb = color
    run.font.size = Pt(size)
    return p


def add_h1(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = NAVY
    run.font.size = Pt(18)
    return p


def add_h2(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = NAVY
    run.font.size = Pt(14)
    return p


def add_h3(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = DARK_GRAY
    run.font.size = Pt(12)
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.color.rgb = DARK_GRAY
    run.font.size = Pt(11)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.color.rgb = DARK_GRAY
    run.font.size = Pt(11)
    return p


def add_image(doc, path, width_inches=6.0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=Inches(width_inches))


def add_divider(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), ACCENT_HEX)
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_page_break(doc):
    doc.add_page_break()


def _add_page_numbers(doc):
    """Add automatic page numbers (centered footer)."""
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run.font.size = Pt(10)
    run.font.color.rgb = DARK_GRAY


def _add_running_header(doc, title):
    """Add running header."""
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(title)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = DARK_GRAY


def cover_page(doc, title, subtitle, author, publisher, edition, year, isbn=None):
    """Render a publishing-grade cover page sized to fit a single page."""
    for _ in range(2):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(publisher.upper())
    run.bold = True
    run.font.color.rgb = ACCENT
    run.font.size = Pt(14)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Professional Reference Series")
    run.italic = True
    run.font.color.rgb = DARK_GRAY
    run.font.size = Pt(11)
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.color.rgb = NAVY
    run.font.size = Pt(28)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(subtitle)
    run.italic = True
    run.font.color.rgb = DARK_GRAY
    run.font.size = Pt(14)
    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(author)
    run.bold = True
    run.font.color.rgb = NAVY
    run.font.size = Pt(16)
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{edition}  |  {year}")
    run.font.color.rgb = DARK_GRAY
    run.font.size = Pt(11)


def copyright_page(doc, title, author, publisher, edition, year, isbn=None):
    """Render a copyright page."""
    for _ in range(8):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.color.rgb = NAVY
    run.font.size = Pt(13)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"by {author}")
    run.font.color.rgb = DARK_GRAY
    run.font.size = Pt(11)
    doc.add_paragraph()
    doc.add_paragraph()
    items = [
        f"Published by {publisher}, San Francisco Bay Area.",
        f"{edition}, {year}.",
        f"Copyright © {year} {publisher}. All rights reserved.",
        "",
        "No part of this work may be reproduced or transmitted in any form or by "
        "any means without the prior written permission of the publisher, except "
        "for brief quotations in critical reviews and certain other non-commercial "
        "uses permitted by copyright law.",
        "",
        "While every precaution has been taken in the preparation of this book, "
        "the publisher and author assume no responsibility for errors or "
        "omissions, or for damages resulting from the use of the information "
        "contained herein.",
        "",
        "Trademarked product names mentioned in this book are the property of "
        "their respective owners. The use of such names is for editorial purposes "
        "only and does not imply endorsement.",
    ]
    if isbn:
        items.insert(3, f"Reference Identifier (RID): {isbn}")
    for item in items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if item:
            run = p.add_run(item)
            run.font.color.rgb = DARK_GRAY
            run.font.size = Pt(10)


def about_author(doc, author, bio):
    """About-the-author section."""
    p = doc.add_paragraph()
    run = p.add_run("About the Author")
    run.bold = True
    run.font.color.rgb = NAVY
    run.font.size = Pt(22)
    add_divider(doc)
    p = doc.add_paragraph()
    run = p.add_run(author)
    run.bold = True
    run.font.color.rgb = NAVY
    run.font.size = Pt(14)
    p = doc.add_paragraph()
    run = p.add_run(bio)
    run.font.color.rgb = DARK_GRAY
    run.font.size = Pt(11)
